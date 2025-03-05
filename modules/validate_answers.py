import faiss
from sentence_transformers import SentenceTransformer
import numpy as np
import re
import os
import csv
from docx import Document  # Import for handling .docx files

# Step 1: Load and Preprocess the Document
def preprocess_document(text):
    """Extract answers and clean text."""
    # Split answers based on patterns like 'Answer 1', 'Answer 2', etc.
    answers = re.split(r'(Answer \d+)', text)
    processed_answers = []

    for i in range(1, len(answers), 2):
        processed_answers.append(answers[i] + ' ' + answers[i + 1].strip())
    return processed_answers

# Step 2: Embed Answers
model = SentenceTransformer('all-MiniLM-L6-v2')  # Load a lightweight embedding model
def embed_answers(answers):
    """Convert a list of answers into embeddings."""
    return model.encode(answers, convert_to_numpy=True)

# Step 3: Index Key Answers
def build_faiss_index(key_embeddings):
    """Build and return a FAISS index from key answer embeddings."""
    dim = key_embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)  # Use L2 distance metric
    index.add(key_embeddings)      # Add key embeddings
    return index

# Step 4: Compare Student Answers and Score
def compute_scores(student_embeddings, key_index):
    """Find the closest matches and compute scores."""
    distances, indices = key_index.search(student_embeddings, 1)  # Get the nearest match
    return distances, indices

def assign_marks(distances, question_marks):
    """Assign marks based on distances and custom question marks."""
    scores = []
    for i, dist in enumerate(distances):
        if dist <= 0.2:  # High similarity threshold
            scores.append(question_marks[i])
        elif dist <= 0.5:  # Medium similarity threshold
            scores.append(question_marks[i] * 0.5)
        else:  # Low similarity threshold
            scores.append(0)
    return scores

# Step 5: Process Multiple Student Papers
def process_student_papers(student_files, key_file, output_file, question_marks):
    """Process and score multiple student papers."""
    # Load and preprocess key answers from docx file
    with open(key_file, 'rb') as f:
        doc = Document(f)
        key_text = ' '.join([para.text for para in doc.paragraphs])
    key_answers = preprocess_document(key_text)
    key_embeddings = embed_answers(key_answers)

    # Build FAISS Index
    faiss_index = build_faiss_index(key_embeddings)

    # Prepare output CSV
    with open(output_file, 'w', newline='') as csvfile:
        csvwriter = csv.writer(csvfile)
        header = ["Student File"] + [f"Answer {i+1}" for i in range(len(key_answers))]
        csvwriter.writerow(header)

        # Process each student file
        for student_file in student_files:
            with open(student_file, 'rb') as f:
                doc = Document(f)
                student_text = ' '.join([para.text for para in doc.paragraphs])
            student_answers = preprocess_document(student_text)
            student_embeddings = embed_answers(student_answers)

            # Compare and score
            distances, _ = compute_scores(student_embeddings, faiss_index)
            scores = assign_marks(distances.flatten(), question_marks)

            # Write to CSV
            csvwriter.writerow([os.path.basename(student_file)] + scores)

# Example Workflow
if __name__ == "__main__":
    # Paths to input and output
    key_file = r"fear_free_validations\218297601003.docx"  # Path to the original answer sheet in .docx format
    student_files = [f"fear_free_validations/21829760100{i}.docx" for i in range(1, 6)]  # 30 student papers in .docx format
    output_file = "student_scores.csv"  # Output CSV for scores

    # Custom marks per question
    question_marks = [5, 8, 10]  # Example: First question = 3 marks, Second = 5 marks, Third = 10 marks

    # Process and score
    process_student_papers(student_files, key_file, output_file, question_marks)

    print(f"Scores saved to {output_file}")
